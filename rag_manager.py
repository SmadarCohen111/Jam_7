from __future__ import annotations

import re
import io
import json
import os
import threading
from pathlib import Path
from typing import Dict, List, Optional, Union, TYPE_CHECKING

import yaml
from bs4 import BeautifulSoup
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_openai import OpenAIEmbeddings
import markdown
import docx
import openpyxl
from PyPDF2 import PdfReader
import hashlib
import contextlib
import csv
import chromadb  # ✅ ADD THIS IMPORT

try:
    from filelock import FileLock
except ImportError:
    FileLock = None

# Optional helper modules
try:
    from helpers.loggers import CustomLogger
except ModuleNotFoundError:
    import logging

    def _get_logger(name: str):
        logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(name)s – %(message)s")
        return logging.getLogger(name)
else:
    def _get_logger(name: str):
        return CustomLogger(logger_name=name).get_logger()

def _default_persist_dir() -> Path:
    return Path(os.getenv("RAG_PERSIST_DIR", "./vector_db")).resolve()

class RAGManager:
    """Singleton RAG Manager using HTTP ChromaDB Client to prevent multi-agent conflicts."""
    
    # Singleton implementation
    _instance = None
    _instance_lock = threading.Lock()
    
    MAX_STORE_SIZE = 10_000

    if TYPE_CHECKING:
        from filelock import FileLock as _FileLockType

    def _hash_bytes_or_text(data) -> str:
        b = data if isinstance(data, bytes) else data.encode("utf-8", "ignore")
        return hashlib.md5(b).hexdigest()

    def _chunk_id(source: str, version: str, idx: int) -> str:
        # short, stable id
        return hashlib.sha256(f"{source}|{version}|{idx}".encode("utf-8")).hexdigest()[:40]

    def __new__(cls, *args, **kwargs):
        """Singleton pattern to ensure only one RAGManager instance exists."""
        if not cls._instance:
            with cls._instance_lock:
                if not cls._instance:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(
        self,
        *,
        chroma_host: str = "localhost",
        chroma_port: int = 9000,
        text_splitter: Optional[RecursiveCharacterTextSplitter] = None,
        lock_timeout: int = 30,
    ) -> None:
        # Prevent re-initialization of singleton
        if hasattr(self, '_initialized') and self._initialized:
            return

        # ✅ CRITICAL CHANGE: Use HTTP Client instead of file-based storage
        self.chroma_host = chroma_host
        self.chroma_port = chroma_port
        
        # Create HTTP client to connect to ChromaDB server
        try:
            self.chroma_client = chromadb.HttpClient(
                host=self.chroma_host, 
                port=self.chroma_port,
            )
            # Test connection
            self.chroma_client.heartbeat()
            self.logger = _get_logger("RAGManager")
            self.logger.info(f"Connected to ChromaDB server at http://{chroma_host}:{chroma_port}")
        except Exception as e:
            self.logger = _get_logger("RAGManager")
            self.logger.error(f"Failed to connect to ChromaDB server: {e}")
            raise ConnectionError(f"Cannot connect to ChromaDB server at {chroma_host}:{chroma_port}")

        self.embedding_fn = OpenAIEmbeddings(model="text-embedding-3-small")
        self.text_splitter = text_splitter or RecursiveCharacterTextSplitter(
            chunk_size=3000,
            chunk_overlap=400,
        )

        # Simplified internal state - HTTP-based Chroma stores
        self._vector_stores: Dict[str, Chroma] = {}
        self._locks: Dict[str, Union["_FileLockType", threading.Lock, None]] = {}
        
        self.lock_timeout: int = lock_timeout
        self.logger.info("RAGManager initialised as SINGLETON with HTTP ChromaDB client")
        
        # Mark as initialized
        self._reset_seen_ids()
        self._initialized = True

    # ────────────────────────────────────────────────────────────────────
    # ✨ NEW – remove vectors whose IDs we did NOT see this run
    # ────────────────────────────────────────────────────────────────────
    def prune_stale_documents(self, store_label: str, valid_ids: set[str]) -> int:
        """
        Delete embeddings from `store_label` that are *not* in valid_ids.
        Returns the number of vectors removed.
        """
        collection = self._ensure_chroma_collection(store_label)
        if collection is None:
            return 0

        try:
            existing = set(collection._collection.get(ids=None, limit=None)["ids"])
            stale_ids = list(existing - valid_ids)
            if stale_ids:
                collection._collection.delete(ids=stale_ids)
            self.logger.info(
                "[%s] Pruned %s stale chunks (remaining %s)",
                store_label, len(stale_ids), len(existing) - len(stale_ids)
            )
            return len(stale_ids)
        except Exception as e:
            self.logger.error("[%s] Could not prune stale docs: %s", store_label, e)
            return 0

    def _validate_store_label(self, store_label: str) -> str:
        """Validate and normalize store label for ChromaDB compatibility."""
        label = store_label.lower().strip()
        # ChromaDB collection names must be 3-63 chars, alphanumeric + hyphens/underscores
        if not re.match(r'^[a-zA-Z0-9_-]{3,63}$', label):
            raise ValueError(f"Invalid store_label '{label}': must be 3-63 chars, alphanumeric/hyphens/underscores only")
        return label

    # ────────────────────────────────────────────────────────────────────────
    # per‑run bookkeeping of IDs we ingested
    # ────────────────────────────────────────────────────────────────────────
    def _reset_seen_ids(self) -> None:
        """Call once per sync run (e.g. at top of rag_sync.py)."""
        self._seen_ids: Dict[str, set[str]] = {}

    def _mark_ids_seen(self, store_label: str, ids: list[str]) -> None:
        """Called internally by _store_documents()."""
        if not hasattr(self, "_seen_ids"):
            self._reset_seen_ids()
        self._seen_ids.setdefault(store_label, set()).update(ids)

    def get_seen_ids(self, store_label: str) -> set[str]:
        """Return the IDs ingested for *this* run of rag_sync.py."""
        return getattr(self, "_seen_ids", {}).get(store_label, set())

    # ===========================================================================
    # Public API – ingestion (same as before)
    # ===========================================================================
    def auto_ingest(
        self,
        knowledge_location: Union[str, List[str]],
        store_label: str,
        *,
        storage_type: str = "local",
    ) -> None:
        """Ingest from one or more local/S3 paths into ChromaDB collection."""
        store_label = self._validate_store_label(store_label)
        self.logger.debug("[%s] Using HTTP ChromaDB backend", store_label)

        if not knowledge_location:
            self.logger.error("No path(s) provided for RAG knowledge ingestion.")
            return

        paths = [knowledge_location] if isinstance(knowledge_location, str) else knowledge_location

        for path in paths:
            if storage_type == "local":
                abs_path = Path(path).expanduser().resolve()
                self.logger.info("[%s] Resolved absolute path to knowledge base: %s", store_label, abs_path)
                self.logger.info("[%s] Ingesting LOCAL path: %s", store_label, path)
                self._auto_ingest_folder(path, store_label)

            elif storage_type == "s3":
                self.logger.info("[%s] Ingesting from S3 prefix: %s", store_label, path)
                self._auto_ingest_s3_prefix(path, store_label)

            else:
                self.logger.error("Unknown storage_type='%s'", storage_type)

    # -------------------------------------------------------------------------
    # Internal – ingestion helpers (same as before)
    # -------------------------------------------------------------------------
    def _auto_ingest_folder(self, folder_path: Union[str, Path], store_label: str) -> None:
        path = Path(folder_path).expanduser().resolve()
        if not path.is_dir():
            self.logger.error("[%s] Path %s is not a directory.", store_label, path)
            return

        for file_path in path.rglob("*"):
            if file_path.is_file():
                self._process_file(file_path, store_label, is_s3=False)

    def _auto_ingest_s3_prefix(self, s3_prefix: str, store_label: str) -> None:
        import boto3

        s3_client = boto3.client(
            "s3",
            aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
            aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
            region_name=os.getenv("AWS_S3_REGION_NAME"),
        )
        bucket = os.getenv("AWS_STORAGE_BUCKET_NAME")
        paginator = s3_client.get_paginator("list_objects_v2")
        for page in paginator.paginate(Bucket=bucket, Prefix=s3_prefix):
            for obj in page.get("Contents", []):
                key = obj["Key"]
                if not key.endswith("/"):
                    self._process_file(key, store_label, is_s3=True)

    def _process_file(self, file_path: Union[str, Path], store_label: str, *, is_s3: bool) -> None:
        """
        Read, parse, split and prepare documents from a single file (local or S3),
        then store them with stable IDs so re-ingests don't duplicate data.
        """
        path_str = str(file_path)
        ext = str(Path(file_path).suffix).lower()
        parser = self._get_parser_for_extension(ext)
        if parser is None:
            self.logger.warning("[%s] Unsupported format %s – skipping %s", store_label, ext, path_str)
            return

        # Read raw bytes/text first (we hash this for a stable "version" of the file)
        try:
            if is_s3:
                raw_content = self._read_s3_file(path_str, ext)
            else:
                raw_content = self._read_local_file(path_str, ext)
        except Exception as exc:
            self.logger.error("[%s] Error reading file %s – %s", store_label, path_str, exc)
            return

        if not raw_content:
            self.logger.warning("[%s] Empty content – skipping %s", store_label, path_str)
            return

        # Compute a stable version hash from the raw content (bytes or text)
        if isinstance(raw_content, bytes):
            version_hash = hashlib.md5(raw_content).hexdigest()
        else:
            version_hash = hashlib.md5(raw_content.encode("utf-8", "ignore")).hexdigest()

        # Parse to plain text for chunking
        try:
            text = parser(path_str, raw_content)
        except Exception as exc:
            self.logger.error("[%s] Parser failed for %s – %s", store_label, path_str, exc)
            return

        if not text or not text.strip():
            self.logger.warning("[%s] Parsed text empty – skipping %s", store_label, path_str)
            return

        chunks = self.text_splitter.split_text(text)
        if not chunks:
            self.logger.debug("[%s] No non-empty chunks for %s", store_label, path_str)
            return

        # Build LangChain Documents and deterministic IDs
        docs: List[Document] = []
        ids: List[str] = []
        base_meta = {
            "source": path_str,
            "title": Path(path_str).stem,
            "is_s3": is_s3,
            "extension": ext,
            "version": version_hash,
        }

        for idx, chunk in enumerate(chunks):
            content = (chunk or "").strip()
            if len(content) <= 5:
                continue
            # Stable ID: sha256(source|version|idx) (shortened for readability)
            stable_id = hashlib.sha256(f"{path_str}|{version_hash}|{idx}".encode("utf-8")).hexdigest()[:40]
            ids.append(stable_id)
            docs.append(Document(page_content=content, metadata={**base_meta, "chunk": idx}))

        if not docs:
            self.logger.debug("[%s] All chunks were trivial for %s", store_label, path_str)
            return

        self._store_documents(docs, store_label, ids=ids)

    
    def debug_collection_status(self, store_label: str) -> Dict[str, any]:
        """Debug method to check collection status."""
        store_label = self._validate_store_label(store_label)
        
        try:
            # Check server connection
            self.chroma_client.heartbeat()
            
            # List all collections
            all_collections = [col.name for col in self.chroma_client.list_collections()]
            
            # Check if our collection exists
            collection_exists = store_label in all_collections
            
            # Get collection details if it exists
            collection_details = None
            if collection_exists:
                try:
                    collection = self.chroma_client.get_collection(store_label)
                    collection_details = {
                        "name": collection.name,
                        "count": collection.count(),
                        "metadata": collection.metadata
                    }
                except Exception as e:
                    collection_details = {"error": str(e)}
            
            return {
                "server_connected": True,
                "all_collections": all_collections,
                "collection_exists": collection_exists,
                "collection_details": collection_details,
                "in_memory_cache": store_label in self._vector_stores
            }
        
        except Exception as e:
            return {
                "server_connected": False,
                "error": str(e)
            }

    def _ensure_chroma_collection(self, store_label: str) -> Optional[Chroma]:
        """Ensure a Chroma collection exists and return it."""
        if store_label in self._vector_stores:
            return self._vector_stores[store_label]

        try:
            # Check if collection already exists
            existing_collections = [col.name for col in self.chroma_client.list_collections()]
            collection_exists = store_label in existing_collections
            
            if collection_exists:
                self.logger.info(f"[{store_label}] Collection already exists, connecting to it")
            else:
                self.logger.info(f"[{store_label}] Creating new collection")
                
            # Get or create collection
            collection = self.chroma_client.get_or_create_collection(name=store_label)
            
            # Create Chroma store instance
            chroma_store = Chroma(
                client=self.chroma_client,
                collection_name=store_label,
                embedding_function=self.embedding_fn
            )
            
            # Verify the collection is accessible
            try:
                count = chroma_store._collection.count()
                self.logger.info(f"[{store_label}] Collection verified with {count} existing documents")
            except Exception as e:
                self.logger.warning(f"[{store_label}] Collection created but cannot verify count: {e}")
            
            self._vector_stores[store_label] = chroma_store
            return chroma_store
            
        except Exception as e:
            self.logger.error(f"[{store_label}] Failed to create/access HTTP Chroma collection: {e}", exc_info=True)
            return None

    def _store_documents(
        self,
        docs: List[Document],
        store_label: str,
        ids: Optional[List[str]] = None
    ) -> None:
        """
        Ingests `docs` into the collection `store_label`.

        * Uses deterministic `ids` to avoid duplicates.
        * Prefers `collection.upsert()` when available (Chroma ≥ 0.4.8).
        * Records every ID that was processed this run so we can later prune
          vectors whose source files were deleted.
        """
        if not docs:
            return

        # Drop trivial docs defensively
        docs = [d for d in docs if d.page_content and len(d.page_content.strip()) > 5]
        if not docs:
            self.logger.debug("[%s] All candidate docs were empty or too small", store_label)
            return

        chroma_store = self._ensure_chroma_collection(store_label)
        if chroma_store is None:
            self.logger.error("[%s] Failed to create or get collection", store_label)
            return

        try:
            initial_count = chroma_store._collection.count()
        except Exception:
            initial_count = None

        try:
            # ---------------------------------------------------------------
            # Preferred path: explicit upsert
            # ---------------------------------------------------------------
            if (
                ids is not None
                and hasattr(chroma_store, "_collection")
                and hasattr(chroma_store._collection, "upsert")
            ):
                texts      = [d.page_content for d in docs]
                metadatas  = [d.metadata      for d in docs]
                embeddings = self.embedding_fn.embed_documents(texts)

                chroma_store._collection.upsert(
                    ids=ids,
                    documents=texts,
                    metadatas=metadatas,
                    embeddings=embeddings,
                )
                self.logger.info(
                    "[%s] Upserted %d documents (deterministic IDs)",
                    store_label, len(docs)
                )
            # ---------------------------------------------------------------
            # Fallback path: LangChain add_documents (may reject duplicates)
            # ---------------------------------------------------------------
            else:
                if ids is not None:
                    chroma_store.add_documents(docs, ids=ids)
                else:
                    chroma_store.add_documents(docs)
                self.logger.info(
                    "[%s] Added %d documents via LangChain API",
                    store_label, len(docs)
                )

            # ---------------------------------------------------------------
            # NEW – remember every ID we ingested this run
            # ---------------------------------------------------------------
            if ids:
                self._mark_ids_seen(store_label, ids)

            # ---------------------------------------------------------------
            # Verify collection count (best effort)
            # ---------------------------------------------------------------
            try:
                final_count = chroma_store._collection.count()
                if initial_count is not None:
                    delta = final_count - initial_count
                    self.logger.info(
                        "[%s] Collection count: %s → %s (Δ=%s)",
                        store_label, initial_count, final_count, delta
                    )
                else:
                    self.logger.info(
                        "[%s] Collection count (post‑write): %s",
                        store_label, final_count
                    )
            except Exception as e:
                self.logger.debug(
                    "[%s] Could not verify count after write: %s",
                    store_label, e
                )

        except Exception as e:
            self.logger.error(
                "[%s] Failed to store documents: %s",
                store_label, e, exc_info=True
            )

    # ✅ UPDATED: Public API methods for HTTP client
    def get_vector_store(self, store_label: str) -> Optional[Chroma]:
        """Return the HTTP Chroma vector store instance for store_label."""
        store_label = self._validate_store_label(store_label)
        return self._vector_stores.get(store_label)

    def reload_chroma_store(self, store_label: str) -> None:
        """Reload a Chroma store from HTTP server."""
        store_label = self._validate_store_label(store_label)
        self.logger.info(f"[{store_label}] Attempting reload from HTTP server")

        try:
            # Remove from memory cache
            if store_label in self._vector_stores:
                del self._vector_stores[store_label]
            
            # Reconnect to HTTP collection
            chroma_store = Chroma(
                client=self.chroma_client,
                collection_name=store_label,
                embedding_function=self.embedding_fn
            )
            
            # Test the connection and get count
            try:
                count = chroma_store._collection.count()
                self._vector_stores[store_label] = chroma_store
                self.logger.info(f"[{store_label}] Reloaded HTTP Chroma store successfully (document count={count})")
            except Exception as count_error:
                self.logger.warning(f"[{store_label}] Collection exists but couldn't get count: {count_error}")
                self._vector_stores[store_label] = chroma_store
                
        except Exception as e:
            self.logger.error(f"[{store_label}] Failed to reload HTTP Chroma store: {e}")

    def get_fresh_vector_store(self, store_label: str) -> Optional[Chroma]:
        """Get a fresh vector store from HTTP server."""
        store_label = self._validate_store_label(store_label)
        self.reload_chroma_store(store_label)
        return self._vector_stores.get(store_label)

    def health_check(self, store_label: str) -> Dict[str, any]:
        """Get health information about an HTTP ChromaDB store."""
        store_label = self._validate_store_label(store_label)
        
        try:
            # Check ChromaDB server health
            self.chroma_client.heartbeat()
            server_healthy = True
        except Exception as e:
            return {
                "status": "server_error",
                "store_label": store_label,
                "error": f"ChromaDB server unreachable: {e}"
            }
        
        store = self.get_vector_store(store_label)
        if not store:
            return {
                "status": "collection_not_found", 
                "store_label": store_label,
                "server_healthy": server_healthy
            }
        
        try:
            count = store._collection.count()
            capacity_percent = round((count / self.MAX_STORE_SIZE) * 100, 1)
            
            return {
                "status": "healthy",
                "store_label": store_label,
                "document_count": count,
                "max_capacity": self.MAX_STORE_SIZE,
                "capacity_used_percent": capacity_percent,
                "server_host": self.chroma_host,
                "server_port": self.chroma_port,
                "server_healthy": server_healthy
            }
        except Exception as e:
            return {
                "status": "error", 
                "store_label": store_label,
                "error": str(e),
                "server_healthy": server_healthy
            }

    # ✅ Keep all your existing parser methods (unchanged)
    def _get_parser_for_extension(self, ext: str):
        return {
            ".txt": self._parse_txt,
            ".md": self._parse_markdown,
            ".markdown": self._parse_markdown,
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".xlsx": self._parse_xlsx,
            ".csv": self._parse_csv,
            ".py": self._parse_code,
            ".js": self._parse_code,
            ".jsx": self._parse_code,
            ".json": self._parse_json,
            ".yaml": self._parse_yaml,
            ".yml": self._parse_yaml,
        }.get(ext)

    def _read_local_file(self, file_path: Union[str, Path], ext: str):
        mode = "rb" if ext in {".pdf", ".docx", ".xlsx"} else "r"
        with open(file_path, mode, encoding=None if mode == "rb" else "utf-8", errors="ignore") as f:
            return f.read()

    def _read_s3_file(self, key: str, ext: str):
        from helpers.aws_s3 import read_s3_text_file, read_s3_binary_file
        if ext in {".pdf", ".docx", ".xlsx"}:
            return read_s3_binary_file(key)
        return read_s3_text_file(key)

    # All your existing parser methods remain the same...
    @staticmethod
    def _parse_txt(file_path, file_content):
        return file_content if isinstance(file_content, str) else file_content.decode("utf-8", errors="ignore")

    def _parse_markdown(self, file_path, file_content):
        text = (
            file_content
            if isinstance(file_content, str)
            else file_content.decode("utf-8", errors="ignore")
        )
        html = markdown.markdown(text, extensions=["tables"])
        soup = BeautifulSoup(html, "html.parser")
        plain = soup.get_text(separator="\n")
        import re
        plain = re.sub(r"\n{3,}", "\n\n", plain)
        return plain.strip()

    @staticmethod
    def _parse_pdf(file_path, file_content: bytes):
        reader = PdfReader(io.BytesIO(file_content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)

    @staticmethod
    def _parse_docx(file_path, file_content: bytes):
        doc = docx.Document(io.BytesIO(file_content))
        return "\n".join(p.text for p in doc.paragraphs)

    @staticmethod
    def _parse_xlsx(file_path, file_content: bytes):
        wb = openpyxl.load_workbook(io.BytesIO(file_content), read_only=True)
        return "\n".join(
            " ".join(str(cell) for cell in row if cell is not None)
            for sheet in wb.worksheets for row in sheet.iter_rows(values_only=True)
        )

    @staticmethod
    def _parse_csv(file_path, file_content):
        text = file_content if isinstance(file_content, str) else file_content.decode("utf-8", errors="ignore")
        reader = csv.reader(text.splitlines())
        return "\n".join(" ".join(row) for row in reader if row)

    @staticmethod
    def _parse_code(file_path, file_content):
        return file_content if isinstance(file_content, str) else file_content.decode("utf-8", errors="ignore")

    @staticmethod
    def _parse_json(file_path, file_content):
        try:
            data = json.loads(file_content if isinstance(file_content, str) else file_content.decode("utf-8", errors="ignore"))
            return json.dumps(data, indent=2)
        except Exception:
            return ""

    @staticmethod
    def _parse_yaml(file_path, file_content):
        try:
            data = yaml.safe_load(file_content if isinstance(file_content, str) else file_content.decode("utf-8", errors="ignore"))
            return yaml.dump(data, sort_keys=False)
        except Exception:
            return ""
